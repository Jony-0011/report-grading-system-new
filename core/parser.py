import docx
from PyPDF2 import PdfReader
from PIL import Image
import io

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    headings = []
    images_count = 0
    
    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        
        if text:
            full_text.append(text)
            
            if 'Heading' in style or style in ['标题 1', '标题 2', '标题 3']:
                headings.append(text)
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            images_count += 1
    
    return {
        'text': '\n'.join(full_text),
        'headings': headings,
        'images_count': images_count,
        'paragraphs': [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    }

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    full_text = []
    headings = []
    images_count = 0
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text.append(text)
    
    text_content = '\n'.join(full_text)
    lines = text_content.split('\n')
    
    for line in lines[:30]:
        stripped = line.strip()
        if stripped and len(stripped) < 50 and stripped.istitle():
            headings.append(stripped)
    
    return {
        'text': text_content,
        'headings': headings[:10],
        'images_count': images_count,
        'paragraphs': [p.strip() for p in lines if p.strip()]
    }

def parse_document(file_path, file_type):
    if file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")